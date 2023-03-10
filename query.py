import openai
import docx
import os
import time
import re
import pyfiglet

openai.api_key = os.environ["OPENAI_API_KEY"]
engine="text-davinci-003"
min_response_lines = 15
max_tokens = 2000

def ask_gpt(prompt):
    prompt += "\n"
    response = openai.Completion.create(
        engine=engine,
        prompt=prompt,
        max_tokens=max_tokens,
        top_p = 1,
        frequency_penalty = 0,
        presence_penalty = 0,
        n=2,
        stop=None,
        temperature=0.8
    )
    answer = response.choices[0].text.strip()
    while len(answer.split("\n")) < min_response_lines:
        prompt = "Explain with Examples " + prompt
        response = openai.Completion.create(
        engine=engine,
        top_p = 1,
        frequency_penalty = 0,
        presence_penalty = 0,
        prompt=prompt,
        max_tokens=max_tokens,
        n=1,
        stop=None,
        temperature=0.8
        )
        answer = answer + "\n\n" + response.choices[0].text.strip()
    return answer

    
ascii_banner = pyfiglet.figlet_format("Chat GPT Query Automation")
print(ascii_banner)

print("\n\nChat GPT Query automation script is running.\n\n")
document = docx.Document()

with open('questions.txt', 'r', encoding='utf-8') as f:
    questions = f.readlines()
    current_chapter = ""
    current_subchapter = ""
    question_number = 0
    first_chapter = True 
    for i in range(len(questions)):
        if "Chapter-" in questions[i]:
            if not first_chapter:
                document.add_page_break()
            else:
                first_chapter = False
            current_chapter = questions[i].strip()
            document.add_heading(current_chapter, level=1)
            current_subchapter = ""
            question_number = 0
            pageBreakFlag = True
        elif "Sub Chapter:" in questions[i]:
            current_subchapter = questions[i].strip()
            document.add_heading(current_subchapter, level=2)
            question_number = 0
            pageBreakFlag = False
        elif "Question:" in questions[i]:
            match = re.search('\$(.*?)\$', questions[i])
            pageBreakFlag = False
            if match:
                question_number += 1
                question = match.group(1)
                question_label = f"{current_chapter} - {current_subchapter} - Question {question_number} - {question}?"
                print("Now Executing " + question_label)
                document.add_heading(question, level=3)
                response = ask_gpt(question)
                document.add_paragraph(f"{response}")


document.save("chat_gpt_response.docx")
print("Chat GPT Query automation script execution completed.")
